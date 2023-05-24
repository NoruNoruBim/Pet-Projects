import os
from copy import deepcopy
from openpyxl.reader.excel import load_workbook
import docx
from tqdm import tqdm
from collections import defaultdict
import re


#   parse SRD .docx document and ICD .xlsx document
class DocsParser:
    # __srd_name                - name of SRD like 'SRD *** 06.09.docx'
    # __icd_name                - name of ICD like '***.xlsx'
    # __layer_ids               - list of layers id (type - str)
    # __ignored_reqs            - requirements which will be ignored
    # __req_to_widgets          - dict, where
    #                               key - srd id '***' (str)
    #                               value - set of widget ids like {'345', '567', ...} from that srd (set<str>)
    # __all_srd_text            - dict, where
    #                               key - srd id (str)
    #                               value - requirement text (str)
    # __ds_to_msg               - dict, where
    #                               key - ds variable name (str)
    #                               value - message name (str)
    def __init__(self, set_all=False):
        input_docs = os.listdir('input')
        self.__srd_name = list(filter(lambda x: x[-5:] == '.docx' and x[0] != '~', input_docs))[0]
        self.__icd_name = list(filter(lambda x: x[-5:] == '.xlsx' and x[0] != '~', input_docs))[0]
        self.__layer_ids = {'1', '2', '3', '4', '5', '6', '7', '8', '9', '10', '13',\
                            '14', '15', '16', '17', '18', '19', '20', '21', '22', '23'}
        with open('input/ignored_reqs.txt', 'r') as f:
            self.__ignored_reqs = set(f.read().split())
        self.__req_to_widgets, self.__all_srd_text = None, None
        self.__ds_to_msg = None
        if set_all:
            self.set_req_to_widgets()
            self.set_ds_to_msg()


    #   extract all text from table without reqursion (requirements tree)
    def __extract_text_from_table(self, table):
        table_text = ''
        for row in table.rows:
            old_cell = None
            for cell in row.cells:
                if old_cell != cell:
                    table_text += '\n' + cell.text
        return table_text
    
    
    #   extract all widget id from text using regular expression
    def __extract_id_from_text(self, text, rude_search=True):
        ids = set()
        if rude_search:
            pattern = r'[0-9]+'
        else:
            pattern = r'\([0-9]*\,? ?[0-9]+\)'
            for i in re.findall(r'\(?\<[\, 0-9]*\>\)?', text):#     ['(<123, 456, 789, ...>)', ...] or just <...>
                elem = i.strip('<>()').replace(',', ' ').split()
                ids.update(set(elem))

        for i in re.findall(pattern, text):#   ['(1, 123)', '(5, 345)', '(654)', ...]
            elem = i.split(',')[-1].strip(' ()')
            if elem not in self.__layer_ids and len(elem) != 0 and '0' != elem[0]:
                ids.add(elem)

        return ids


    #   extract dict like {srd : {widget ids}, ...}
    def __extract_req_to_widgets(self):
        doc = docx.Document(os.getcwd() + r'/input/' + self.__srd_name)
        
        tq = tqdm(total=len(doc.tables))
        req_to_widgets = defaultdict(set)
        flag = False
        req = ''
        all_text = defaultdict(str)
        for table in doc.tables:#       for each table in SRD
            tq.update(1)
            '''
            github
            here is removed code
            '''
        tq.close()       

        return req_to_widgets, all_text
    
    
    #   make trace between ds variables and message names using ICD
    def __extract_ds_to_msg(self):
        wb = load_workbook(os.getcwd() + '/input/' + self.__icd_name)
        sheet = wb['Data_Definition']
        ds_to_msg = {}
        for i in range(3, len(sheet['K']) + 1):
            if sheet.cell(i, 11).value not in ds_to_msg:
                ds_to_msg[sheet.cell(i, 11).value] = sheet.cell(i, 6).value
        
        return ds_to_msg


    #   set-method
    def set_req_to_widgets(self):
        if self.__req_to_widgets == None or self.__all_srd_text == None:
            self.__req_to_widgets, self.__all_srd_text = self.__extract_req_to_widgets()
            actual_req_to_widget_keys = list(filter(lambda x: x not in self.__ignored_reqs, self.__req_to_widgets))
            actual_all_srd_text_keys = list(filter(lambda x: x not in self.__ignored_reqs, self.__all_srd_text))
            '''
            github
            here is removed code
            '''


    #   set-method
    def set_ds_to_msg(self):
        if self.__ds_to_msg == None:
            self.__ds_to_msg = self.__extract_ds_to_msg()


    #   get-method
    def get_req_to_widgets(self):
        return deepcopy(self.__req_to_widgets)


    #   get-method
    def get_all_srd_text(self):
        return self.__all_srd_text


    #   get-method
    def get_ds_to_msg(self):
        return deepcopy(self.__ds_to_msg)


    # key   - which field we want to print
    # limit - max number of elements which will be printed
    def print_data(self, key, limit=None, to_file=False, filename='file.txt'):#    print-method
        if key == 'req_to_widgets':
            '''
            github
            here is removed code
            '''
        elif key == 'all_srd_text':
            '''
            github
            here is removed code
            '''
        elif key == 'ds_to_msg':
            '''
            github
            here is removed code
            '''
        else:
            print("- Невозможно вывести данные. Используйте следующие ключи: 'req_to_widgets', 'all_srd_text', 'ds_to_msg'\n")
        
        if to_file:
            with open(filename, 'w', encoding='utf-8') as f:
                f.write(s)
        else:
            print(s)
